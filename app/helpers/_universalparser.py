import json
from typing import List, Optional

from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PDFMinerLoader
from langchain.docstore.document import Document as LangchainDocument
import magic

from ._textcleaner import TextCleaner
from app.schemas.files import JsonFile


class UniversalParser:
    DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    PDF_TYPE = "application/pdf"
    JSON_TYPE = "application/json"
    TXT_TYPE = "text/plain"
    CSV_TYPE = "text/csv"  # separators should be = ";"

    SUPPORTED_FILE_TYPES = [DOCX_TYPE, PDF_TYPE, JSON_TYPE]

    def __init__(self):
        """
        Initializes the class instance.

        Attributes:
            cleaner (TextCleaner): An instance of TextCleaner used for cleaning text.
        """
        self.cleaner = TextCleaner()
        pass

    def parse_and_chunk(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int,
        chunk_min_size: int,
    ):
        """
        Parses a file and splits it into text chunks based on the file type.

        Args:
            file_path (str): Path to the file to be processed.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of characters overlapping between chunks.
            chunk_min_size (int): Minimum size of a chunk to be considered valid.

        Returns:
            list: List of Langchain documents, where each document corresponds to a text chunk.

        Raises:
            NotImplementedError: If the file type is not supported.
        """
        file_type = magic.from_file(file_path, mime=True)

        if file_type == self.TXT_TYPE:
            # In the case the json file is stored as text/plain instead of application/json
            with open(file_path, "r") as file:
                content = file.read()
                try:
                    data = json.loads(content)
                    file_type = self.JSON_TYPE
                except json.JSONDecodeError:
                    pass

        if file_type not in self.SUPPORTED_FILE_TYPES:
            file_type = "unknown"

        if file_type == self.PDF_TYPE:
            chunks = self._pdf_to_chunks(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                chunk_min_size=chunk_min_size,
            )
        elif file_type == self.DOCX_TYPE:
            chunks = self._docx_to_chunks(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                chunk_min_size=chunk_min_size,
            )

        elif file_type == self.JSON_TYPE:
            chunks = self._json_to_chunks(
                file_path=file_path,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                chunk_min_size=chunk_min_size,
            )

        else:
            raise NotImplementedError(f"Unsupported input file format ({file_path}): {file_type}")

        return chunks

    ## Parser and chunking functions

    def _pdf_to_chunks(
        self, file_path: str, chunk_size: int, chunk_overlap: int, chunk_min_size: int
    ) -> List[LangchainDocument]:
        """
        Parse a PDF file and converts it into a list of text chunks.

        Args:
            file_path (str): Path to the PDF file to be processed.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of characters overlapping between chunks.
            chunk_min_size (int): Minimum size of a chunk to be considered valid.

        Returns:
            list: List of Langchain documents, where each document corresponds to a text chunk.
        """

        loader = PDFMinerLoader(file_path)
        doc = loader.load()

        chunks = []
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=["\n\n", "\n"],
        )
        # Splitting text because too long
        splitted_text = text_splitter.split_text(doc[0].page_content)

        for k, text in enumerate(splitted_text):
            if chunk_min_size:
                if len(text) > chunk_min_size:  # We avoid meaningless little chunks
                    chunk = LangchainDocument(
                        page_content=self.cleaner.clean_string(text),
                        metadata={
                            "file_id": file_path.split("/")[-1],
                        },
                    )
            else:
                chunk = LangchainDocument(
                    page_content=self.cleaner.clean_string(text),
                    metadata={
                        "file_id": file_path.split("/")[-1],
                    },
                )
            chunks.append(chunk)

        return chunks  # List of langchain documents

    def _docx_to_chunks(
        self,
        file_path: str,
        chunk_size: int,
        chunk_overlap: int,
        chunk_min_size: int,
    ) -> List[LangchainDocument]:
        """
        Parse a DOCX file and converts it into a list of text chunks.

        Args:
            file_path (str): Path to the DOCX file to be processed.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of characters overlapping between chunks.
            chunk_min_size (int): Minimum size of a chunk to be considered valid.

        Returns:
            list: List of Langchain documents, where each document corresponds to a text chunk.
        """
        documents = []

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
        )

        doc = Document(file_path)
        title = None
        text_chunks = []

        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith("Heading"):
                if title:
                    full_text = "\n".join([p.text for p in text_chunks])
                    splitted_text = text_splitter.split_text(full_text)
                    # Adding previous subpart to result

                    if splitted_text:
                        for k, text in enumerate(splitted_text):
                            if chunk_min_size:
                                if len(text) > chunk_min_size:  # We avoid meaningless little chunks
                                    chunk = LangchainDocument(
                                        page_content=self.cleaner.clean_string(text),
                                        metadata={
                                            "file_id": file_path.split("/")[-1],
                                            "title": title,
                                        },
                                    )
                            else:
                                chunk = LangchainDocument(
                                    page_content=self.cleaner.clean_string(text),
                                    metadata={
                                        "file_id": file_path.split("/")[-1],
                                        "title": title,
                                    },
                                )
                            documents.append(chunk)
                # Updating title for new subpart
                title = paragraph.text.strip()
                text_chunks = []
            else:
                text_chunks.append(paragraph)

        # Adding the last subpart
        if title:
            full_text = "\n".join([p.text for p in text_chunks])
            splitted_text = text_splitter.split_text(full_text)

            if splitted_text:
                for k, text in enumerate(splitted_text):
                    if chunk_min_size:
                        if len(text) > chunk_min_size:  # We avoid meaningless little chunks
                            chunk = LangchainDocument(
                                page_content=self.cleaner.clean_string(text),
                                metadata={
                                    "file_id": file_path.split("/")[-1],
                                    "title": title,
                                },
                            )
                    else:
                        chunk = LangchainDocument(
                            page_content=self.cleaner.clean_string(text),
                            metadata={
                                "file_id": file_path.split("/")[-1],
                                "title": title,
                            },
                        )
                    documents.append(chunk)

        elif text_chunks:
            full_text = "\n".join([p.text for p in text_chunks])
            splitted_text = text_splitter.split_text(full_text)

            for k, text in enumerate(splitted_text):
                if chunk_min_size:
                    if len(text) > chunk_min_size:  # We avoid meaningless little chunks
                        chunk = LangchainDocument(
                            page_content=self.cleaner.clean_string(text),
                            metadata={
                                "file_id": file_path.split("/")[-1],
                                "title": title,
                            },
                        )
                else:
                    chunk = LangchainDocument(
                        page_content=self.cleaner.clean_string(text),
                        metadata={
                            "file_id": file_path.split("/")[-1],
                            "title": title,
                        },
                    )
                documents.append(chunk)

        return documents

    def _json_to_chunks(
        self,
        file_path: str,
        chunk_size: Optional[int],
        chunk_overlap: Optional[int],
        chunk_min_size: Optional[int],
    ) -> List[LangchainDocument]:
        """
        Converts a JSON file into a list of chunks.

        Args:
            file_path (str): Path to the JSON file to be processed.
            chunk_size (int): Maximum size of each text chunk.
            chunk_overlap (int): Number of characters overlapping between chunks.
            chunk_min_size (int): Minimum size of a chunk to be considered valid.

        Returns:
            list: List of Langchain documents, where each document corresponds to a text chunk.
        """

        with open(file_path, "r") as file:
            data = json.load(file)
            data = JsonFile(**data)  # Validate the JSON file

        (file_path.split("/")[-1],)

        chunks = []
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=["\n"],
        )

        for i, document in enumerate(data.documents):
            if document.metadata:
                document.metadata["file_id"] = file_path.split("/")[-1]
            else:
                document.metadata = {"file_id": file_path.split("/")[-1]}

            splitted_text = text_splitter.split_text(document.text)
            for k, text in enumerate(splitted_text):
                if (
                    chunk_min_size and len(text) < chunk_min_size
                ):  # We avoid meaningless little chunks
                    continue

                chunk = LangchainDocument(
                    page_content=self.cleaner.clean_string(text),
                    metadata=document.metadata,
                )
                chunks.append(chunk)

        return chunks
